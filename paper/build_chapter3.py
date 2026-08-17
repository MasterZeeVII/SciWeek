# -*- coding: utf-8 -*-
"""Build บทที่ 3 (การวิเคราะห์และออกแบบระบบ) as a .docx.

Formatting follows ref/project1-v33003.docx: TH SarabunPSK, 16pt body,
18pt chapter heading. Diagrams come from paper/diagrams/ (make_diagrams.py).
"""
import os

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(HERE, "diagrams")
OUT = os.path.join(HERE, "บทที่3_การวิเคราะห์และออกแบบระบบ.docx")

FONT = "TH SarabunPSK"


def style_run(run, size=16, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)  # Thai is complex script
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = rPr.makeelement(qn("w:szCs"), {})
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size * 2))
    if bold:
        bCs = rPr.find(qn("w:bCs"))
        if bCs is None:
            bCs = rPr.makeelement(qn("w:bCs"), {})
            rPr.append(bCs)
        bCs.set(qn("w:val"), "1")


def para(doc, text="", size=16, bold=False, align=None, indent=None, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if text:
        style_run(p.add_run(text), size=size, bold=bold)
    return p


def figure(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    num, rest = caption.split(" ", 1)[0] + " " + caption.split(" ", 1)[1].split(" ", 1)[0], caption.split(" ", 2)[2]
    style_run(cap.add_run(num + " "), bold=True)
    style_run(cap.add_run(rest))


def landscape(doc):
    """Start a new landscape A4 page; returns the section."""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width = Cm(29.7)
    s.page_height = Cm(21.0)
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.0)
    return s


def portrait(doc):
    """Return to portrait A4 with thesis margins."""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(3.81)
    s.left_margin = Cm(3.81)
    s.right_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    return s


def event_table(doc, caption, rows):
    """rows: list of (no, event, response)."""
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    num = " ".join(caption.split(" ")[:2])
    rest = caption.split(" ", 2)[2]
    style_run(cap.add_run(num + " "), bold=True)
    style_run(cap.add_run(rest))

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.4), Cm(7.6), Cm(8.8)]
    headers = ["ลำดับ", "เหตุการณ์ (Event)", "การตอบสนองของระบบ (Response)"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(h), size=14, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths[c]
            p = cell.paragraphs[0]
            if c == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(str(val)), size=14)
    para(doc, "", space_after=6)


EVENTS = [
    (1, "ผู้ดูแลระบบสร้างการแข่งขันประจำปีและกำหนดรุ่นการแข่งขัน",
        "บันทึกการแข่งขัน (จัดได้มากกว่าหนึ่งครั้งต่อปีโดยแยกด้วยครั้งที่ของการแข่งขัน) "
        "สร้างรุ่นมัธยมศึกษาตอนต้น/ตอนปลาย และเปิดใช้งานการแข่งขันรายการนั้น"),
    (2, "ผู้ดูแลระบบลงทะเบียนโรงเรียน ทีม และรายชื่อผู้เล่น",
        "บันทึกทีมเข้าสังกัดรุ่นการแข่งขันและโรงเรียน พร้อมรายชื่อสมาชิกทีม"),
    (3, "ผู้ดูแลระบบกำหนดรอบการแข่งขัน (Best of) และสั่งสร้างสายการแข่งขัน",
        "สุ่มจับสายแบบแพ้คัดออก (Single Elimination) สร้างรอบ คู่แข่งขัน และเกมของรอบแรก"),
    (4, "เจ้าหน้าที่สนามอัปโหลดภาพถ่ายหน้าจอสรุปผลพร้อมระบุทีมชนะและตำแหน่งคะแนน (ROI)",
        "สกัดคะแนนด้วยเทคโนโลยี OCR บันทึกภาพหลักฐาน และแสดงรายการรอตรวจสอบ"),
    (5, "ผู้ตรวจสอบผลยืนยันผลเกม",
        "บันทึกคะแนนจริง ปรับสถานะคู่แข่งขัน และเลื่อนทีมชนะเข้าสู่คู่แข่งขันรอบถัดไปโดยอัตโนมัติ"),
    (6, "ผู้ตรวจสอบผลปฏิเสธผลเกม",
        "บันทึกเหตุผลการปฏิเสธ และเปิดให้เจ้าหน้าที่สนามอัปโหลดภาพหลักฐานใหม่"),
    (7, "มีการแก้ไขผลเกมของคู่แข่งขันที่ตัดสินผลไปแล้ว",
        "ล้างผลของคู่แข่งขันรอบถัดไปที่อ้างอิงผลเดิมโดยอัตโนมัติ เพื่อป้องกันข้อมูลขัดแย้ง"),
    (8, "ผู้ใช้งานทั่วไปเรียกดูผลการแข่งขันผ่านเว็บไซต์",
        "แสดงสายการแข่งขัน ผลคะแนน และสถิติ เฉพาะเกมที่ผ่านการยืนยันแล้วเท่านั้น"),
]


def data_table(doc, caption, rows):
    """rows: list of (no, field, type, description, key)."""
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    num = " ".join(caption.split(" ")[:2])
    rest = caption.split(" ", 2)[2]
    style_run(cap.add_run(num + " "), bold=True)
    style_run(cap.add_run(rest))

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.4), Cm(4.2), Cm(3.4), Cm(6.2), Cm(2.6)]
    headers = ["ลำดับ", "ชื่อฟิลด์", "ชนิดข้อมูล", "คำอธิบาย", "คีย์"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(h), size=14, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths[c]
            p = cell.paragraphs[0]
            if c in (0, 4):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(str(val)), size=14)
    para(doc, "", space_after=6)


DD = [
    ("ตารางที่ 3.2 พจนานุกรมข้อมูลตารางผู้ใช้งาน (users)", [
        (1, "id", "INT UNSIGNED", "รหัสผู้ใช้งาน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "username", "VARCHAR(50)", "ชื่อบัญชีผู้ใช้งาน (ห้ามซ้ำ)", "UQ"),
        (3, "password_hash", "VARCHAR(255)", "รหัสผ่านที่ผ่านการเข้ารหัสแบบแฮช", ""),
        (4, "role", "ENUM", "บทบาทผู้ใช้งาน (ADMIN / MONITOR / FIELD_STAFF)", ""),
        (5, "is_active", "TINYINT(1)", "สถานะการใช้งานบัญชี (ค่าเริ่มต้น 1)", ""),
        (6, "created_at", "DATETIME", "วันเวลาที่สร้างบัญชี", ""),
    ]),
    ("ตารางที่ 3.3 พจนานุกรมข้อมูลตารางโรงเรียน (schools)", [
        (1, "id", "INT UNSIGNED", "รหัสโรงเรียน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "name", "VARCHAR(100)", "ชื่อโรงเรียน (ห้ามซ้ำ)", "UQ"),
    ]),
    ("ตารางที่ 3.4 พจนานุกรมข้อมูลตารางการแข่งขัน (tournaments)", [
        (1, "id", "INT UNSIGNED", "รหัสการแข่งขัน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "name", "VARCHAR(100)", "ชื่อรายการแข่งขัน", ""),
        (3, "year", "SMALLINT UNSIGNED", "ปี พ.ศ. ของการแข่งขัน (ห้ามซ้ำร่วมกับ season)", "UQ"),
        (4, "season", "SMALLINT UNSIGNED", "ครั้งที่ของการแข่งขันภายในปีเดียวกัน (ค่าเริ่มต้น 1, ห้ามซ้ำร่วมกับ year)", "UQ"),
        (5, "is_active", "TINYINT(1)", "สถานะการแข่งขันที่เปิดใช้งาน (เปิดได้ครั้งละ 1 รายการ)", ""),
        (6, "created_at", "DATETIME", "วันเวลาที่สร้างรายการแข่งขัน", ""),
    ]),
    ("ตารางที่ 3.5 พจนานุกรมข้อมูลตารางรุ่นการแข่งขัน (divisions)", [
        (1, "id", "INT UNSIGNED", "รหัสรุ่นการแข่งขัน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "tournament_id", "INT UNSIGNED", "รหัสการแข่งขันที่สังกัด (อ้างอิง tournaments)", "FK"),
        (3, "level", "ENUM", "ระดับรุ่น (JUNIOR = มัธยมต้น / SENIOR = มัธยมปลาย)", ""),
        (4, "max_teams", "TINYINT UNSIGNED", "จำนวนทีมสูงสุดของรุ่น (ค่าเริ่มต้น 32)", ""),
        (5, "default_best_of", "TINYINT UNSIGNED", "รูปแบบ Best of เริ่มต้น (ค่าเริ่มต้น 3)", ""),
    ]),
    ("ตารางที่ 3.6 พจนานุกรมข้อมูลตารางรอบการแข่งขัน (rounds)", [
        (1, "id", "INT UNSIGNED", "รหัสรอบการแข่งขัน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "division_id", "INT UNSIGNED", "รหัสรุ่นการแข่งขันที่สังกัด (อ้างอิง divisions)", "FK"),
        (3, "round_number", "TINYINT UNSIGNED", "ลำดับรอบการแข่งขันภายในรุ่น", ""),
        (4, "round_name", "VARCHAR(50)", "ชื่อรอบการแข่งขัน เช่น รอบรองชนะเลิศ", ""),
        (5, "best_of", "TINYINT UNSIGNED", "จำนวนเกมสูงสุดของแต่ละคู่ในรอบ (BO3/BO5/BO7)", ""),
    ]),
    ("ตารางที่ 3.7 พจนานุกรมข้อมูลตารางทีม (teams)", [
        (1, "id", "INT UNSIGNED", "รหัสทีม (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "division_id", "INT UNSIGNED", "รหัสรุ่นการแข่งขันที่สังกัด (อ้างอิง divisions)", "FK"),
        (3, "school_id", "INT UNSIGNED", "รหัสโรงเรียนต้นสังกัด (อ้างอิง schools)", "FK"),
        (4, "team_number", "TINYINT UNSIGNED", "ลำดับทีมของโรงเรียนภายในรุ่น (กรณีส่งหลายทีม)", ""),
    ]),
    ("ตารางที่ 3.8 พจนานุกรมข้อมูลตารางสมาชิกทีม (team_members)", [
        (1, "id", "INT UNSIGNED", "รหัสสมาชิกทีม (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "team_id", "INT UNSIGNED", "รหัสทีมที่สังกัด (อ้างอิง teams)", "FK"),
        (3, "full_name", "VARCHAR(100)", "ชื่อ-นามสกุลผู้เล่น", ""),
        (4, "in_game_name", "VARCHAR(100)", "ชื่อภายในเกม (ระบุหรือไม่ก็ได้)", ""),
    ]),
    ("ตารางที่ 3.9 พจนานุกรมข้อมูลตารางคู่แข่งขัน (matches)", [
        (1, "id", "INT UNSIGNED", "รหัสคู่แข่งขัน (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "round_id", "INT UNSIGNED", "รหัสรอบการแข่งขันที่สังกัด (อ้างอิง rounds)", "FK"),
        (3, "match_number", "TINYINT UNSIGNED", "ลำดับคู่แข่งขันภายในรอบ", ""),
        (4, "next_match_id", "INT UNSIGNED", "รหัสคู่แข่งขันถัดไปที่ทีมชนะจะเลื่อนไป (อ้างอิง matches)", "FK"),
        (5, "status", "ENUM", "สถานะคู่แข่งขัน (PENDING / IN_PROGRESS / COMPLETED)", ""),
        (6, "scheduled_time", "DATETIME", "กำหนดเวลาแข่งขัน (ระบุหรือไม่ก็ได้)", ""),
    ]),
    ("ตารางที่ 3.10 พจนานุกรมข้อมูลตารางเกมการแข่งขัน (match_games)", [
        (1, "id", "INT UNSIGNED", "รหัสเกม (เพิ่มค่าอัตโนมัติ)", "PK"),
        (2, "match_id", "INT UNSIGNED", "รหัสคู่แข่งขันที่สังกัด (อ้างอิง matches)", "FK"),
        (3, "team1_id", "INT UNSIGNED", "รหัสทีมที่ 1 ของเกม (อ้างอิง teams)", "FK"),
        (4, "team2_id", "INT UNSIGNED", "รหัสทีมที่ 2 ของเกม (อ้างอิง teams)", "FK"),
        (5, "game_number", "TINYINT UNSIGNED", "ลำดับเกมภายในคู่แข่งขัน (เกมที่ 1, 2, ...)", ""),
        (6, "winner_team_id", "INT UNSIGNED", "รหัสทีมที่ชนะเกมนี้ (อ้างอิง teams)", "FK"),
        (7, "kill_team1", "SMALLINT UNSIGNED", "คะแนน Kill ของทีมที่ 1 (บันทึกเมื่อยืนยันผลแล้ว)", ""),
        (8, "kill_team2", "SMALLINT UNSIGNED", "คะแนน Kill ของทีมที่ 2 (บันทึกเมื่อยืนยันผลแล้ว)", ""),
        (9, "image_path", "VARCHAR(255)", "ตำแหน่งไฟล์ภาพหลักฐานหน้าจอสรุปผล", ""),
        (10, "uploaded_by_id", "INT UNSIGNED", "รหัสผู้ใช้งานที่อัปโหลดภาพ (อ้างอิง users)", "FK"),
        (11, "uploaded_at", "DATETIME", "วันเวลาที่อัปโหลดภาพ", ""),
        (12, "ocr_kill_team1", "SMALLINT UNSIGNED", "คะแนน Kill ทีมที่ 1 ที่อ่านได้จาก OCR", ""),
        (13, "ocr_kill_team2", "SMALLINT UNSIGNED", "คะแนน Kill ทีมที่ 2 ที่อ่านได้จาก OCR", ""),
        (14, "raw_ocr_json", "JSON", "ผลลัพธ์ดิบจากการสกัดข้อมูลด้วย OCR", ""),
        (15, "ocr_status", "ENUM", "สถานะการตรวจสอบ (PENDING / UPLOADED / OCR_DONE / VERIFIED / REJECTED)", ""),
        (16, "verified_by_id", "INT UNSIGNED", "รหัสผู้ใช้งานที่ยืนยันผล (อ้างอิง users)", "FK"),
        (17, "verified_at", "DATETIME", "วันเวลาที่ยืนยันผล", ""),
        (18, "reject_reason", "VARCHAR(255)", "เหตุผลกรณีปฏิเสธผลการสแกน", ""),
    ]),
]


def build():
    doc = Document()
    # page setup: A4, thesis margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.81)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    para(doc, "บทที่ 3", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "การวิเคราะห์และออกแบบระบบ", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    para(doc, "ในการพัฒนาระบบจัดการผลการแข่งขันกีฬาอีสปอร์ต งานสัปดาห์วิทยาศาสตร์แห่งชาติ "
              "มหาวิทยาลัยราชภัฏนครสวรรค์ ผู้จัดทำได้ดำเนินการวิเคราะห์ความต้องการของระบบจากกระบวนการ"
              "จัดการแข่งขันเดิม แล้วจึงออกแบบระบบใหม่ให้ครอบคลุมการทำงานของผู้ใช้งานทุกบทบาท "
              "โดยนำเสนอผลการวิเคราะห์และออกแบบด้วยแผนภาพการแบ่งฟังก์ชันการทำงาน (Functional "
              "Decomposition Diagram) รายการเหตุการณ์และการตอบสนอง (Event–Response List) "
              "แผนภาพกระแสข้อมูล (Data Flow Diagram) ซึ่งเริ่มต้นจากระดับ 0 หรือแผนภาพบริบท "
              "(Context Diagram) แล้วจึงแตกรายละเอียดเป็นระดับ 1 ตามลำดับ แผนภาพความสัมพันธ์ของข้อมูล "
              "(Entity Relationship Diagram) และพจนานุกรมข้อมูล (Data Dictionary) ดังรายละเอียดต่อไปนี้",
         indent=1.25, space_after=12)

    # 3.1
    para(doc, "3.1 การวิเคราะห์ระบบ (System Analysis)", bold=True, space_after=6)

    para(doc, "3.1.1 แผนภาพการแบ่งฟังก์ชันการทำงาน (Functional Decomposition Diagram)",
         bold=True, indent=1.25, space_after=6)
    para(doc, "จากการวิเคราะห์ความต้องการของผู้ใช้งาน สามารถแบ่งฟังก์ชันการทำงานหลักของระบบออกเป็น "
              "7 ฟังก์ชัน ได้แก่ (1) การจัดการผู้ใช้งานและสิทธิ์ (2) การจัดการการแข่งขันประจำปี "
              "(3) การจัดการทีมผู้เข้าแข่งขัน (4) การจัดสายการแข่งขันแบบแพ้คัดออก (Single Elimination) "
              "(5) การบันทึกผลการแข่งขันด้วยเทคโนโลยี OCR (6) การตรวจสอบและยืนยันผล และ "
              "(7) การแสดงผลการแข่งขันสาธารณะ ดังภาพที่ 3.1",
         indent=1.25, space_after=6)
    landscape(doc)
    figure(doc, os.path.join(DIAGRAMS, "fig_3_1_decomposition.png"),
           "ภาพที่ 3.1 แผนภาพการแบ่งฟังก์ชันการทำงานของระบบ", width=9.8)
    portrait(doc)

    para(doc, "3.1.2 รายการเหตุการณ์และการตอบสนอง (Event–Response List)",
         bold=True, indent=1.25, space_after=6)
    para(doc, "ก่อนการสร้างแผนภาพกระแสข้อมูล ผู้จัดทำได้รวบรวมเหตุการณ์ (Event) ที่เกิดขึ้นจาก"
              "ผู้เกี่ยวข้องภายนอก และการตอบสนอง (Response) ที่ระบบต้องดำเนินการต่อเหตุการณ์นั้น "
              "เพื่อใช้เป็นข้อมูลตั้งต้นในการกำหนดกระบวนการและกระแสข้อมูลของระบบ ดังตารางที่ 3.1",
         indent=1.25, space_after=8)
    event_table(doc, "ตารางที่ 3.1 รายการเหตุการณ์และการตอบสนองของระบบ", EVENTS)

    para(doc, "3.1.3 แผนภาพบริบท (Context Diagram) หรือแผนภาพกระแสข้อมูล ระดับ 0 (DFD Level 0)",
         bold=True, indent=1.25, space_after=6)
    para(doc, "การสร้างแผนภาพกระแสข้อมูลเริ่มต้นจากระดับ 0 หรือแผนภาพบริบท ซึ่งเป็นแผนภาพภาพรวม"
              "ระดับสูงสุดที่แสดงระบบทั้งหมดเป็นกระบวนการเดียว (หมายเลข 0) พร้อมการไหลของข้อมูล"
              "ระหว่างระบบกับผู้เกี่ยวข้องภายนอก "
              "(External Entity) ทั้ง 4 กลุ่ม ได้แก่ ผู้ดูแลระบบ (Admin) ทำหน้าที่กำหนดข้อมูลการแข่งขัน "
              "ทีม และสายการแข่งขัน เจ้าหน้าที่สนาม (Field Staff) ทำหน้าที่ถ่ายภาพหน้าจอสรุปผลการแข่งขัน"
              "พร้อมระบุทีมชนะส่งเข้าสู่ระบบ ผู้ตรวจสอบผล (Monitor) ทำหน้าที่ตรวจสอบคะแนนที่ระบบ OCR "
              "อ่านได้เทียบกับภาพหลักฐานก่อนยืนยันการบันทึก และผู้ใช้งานทั่วไป (Guest) ที่สามารถเรียกดู"
              "ผลการแข่งขันที่ยืนยันแล้วผ่านเว็บไซต์โดยไม่ต้องเข้าสู่ระบบ ดังภาพที่ 3.2",
         indent=1.25, space_after=6)
    figure(doc, os.path.join(DIAGRAMS, "fig_3_2_context.png"),
           "ภาพที่ 3.2 แผนภาพบริบทของระบบ (Context Diagram / DFD Level 0)", width=5.9)

    para(doc, "3.1.4 แผนภาพกระแสข้อมูล ระดับ 1 (Data Flow Diagram Level 1)",
         bold=True, indent=1.25, space_after=6)
    para(doc, "จากแผนภาพบริบทในภาพที่ 3.2 สามารถแตกรายละเอียดลงเป็นแผนภาพกระแสข้อมูลระดับ 1 "
              "ซึ่งแสดงกระบวนการหลักของระบบจำนวน 7 กระบวนการ "
              "พร้อมแหล่งจัดเก็บข้อมูล (Data Store) จำนวน 9 แหล่ง ซึ่งสอดคล้องกับตารางในฐานข้อมูล "
              "โดยกระบวนการที่ 1 ถึง 4 เป็นงานเตรียมการแข่งขันของผู้ดูแลระบบ กระบวนการที่ 5 "
              "เป็นการบันทึกผลจากภาพถ่ายหน้าจอด้วยเทคโนโลยี OCR โดยเจ้าหน้าที่สนาม กระบวนการที่ 6 "
              "เป็นการตรวจสอบและยืนยันผลโดยผู้ตรวจสอบ ซึ่งเมื่อยืนยันแล้วระบบจะปรับสถานะคู่แข่งขัน"
              "และเลื่อนทีมชนะเข้าสู่รอบถัดไปโดยอัตโนมัติ และกระบวนการที่ 7 เป็นการแสดงผลการแข่งขัน"
              "ต่อสาธารณะ ดังภาพที่ 3.3",
         indent=1.25, space_after=6)
    figure(doc, os.path.join(DIAGRAMS, "fig_3_3_dfd1.png"),
           "ภาพที่ 3.3 แผนภาพกระแสข้อมูล ระดับที่ 1 (Data Flow Diagram Level 1)", width=5.4)

    para(doc, "3.1.5 แผนภาพกระแสข้อมูลระดับที่ 1 ของแต่ละกระบวนการ (Data Flow Diagram Level 1 "
              "รายกระบวนการ)", bold=True, indent=1.25, space_after=6)
    para(doc, "จากกระบวนการหลักทั้ง 7 กระบวนการในภาพที่ 3.3 ผู้จัดทำได้แตกรายละเอียดกระบวนการ"
              "แต่ละกระบวนการลงเป็นกระบวนการย่อยตามแผนภาพการแบ่งฟังก์ชันการทำงานในภาพที่ 3.1 "
              "พร้อมระบุกระแสข้อมูลที่ไหลเข้า-ออกจากแหล่งจัดเก็บข้อมูลจริงของแต่ละกระบวนการย่อย "
              "ดังภาพที่ 3.4 ถึงภาพที่ 3.10",
         indent=1.25, space_after=6)
    for n, (fname, caption) in enumerate([
        ("fig_3_4_dfd_p1.png", "ภาพที่ 3.4 DFD ระดับที่ 1 ของกระบวนการที่ 1 จัดการผู้ใช้งานและสิทธิ์"),
        ("fig_3_5_dfd_p2.png", "ภาพที่ 3.5 DFD ระดับที่ 1 ของกระบวนการที่ 2 จัดการการแข่งขันประจำปี"),
        ("fig_3_6_dfd_p3.png", "ภาพที่ 3.6 DFD ระดับที่ 1 ของกระบวนการที่ 3 จัดการทีมผู้เข้าแข่งขัน"),
        ("fig_3_7_dfd_p4.png", "ภาพที่ 3.7 DFD ระดับที่ 1 ของกระบวนการที่ 4 จัดสายการแข่งขัน"),
        ("fig_3_8_dfd_p5.png", "ภาพที่ 3.8 DFD ระดับที่ 1 ของกระบวนการที่ 5 บันทึกผลการแข่งขันด้วย OCR"),
        ("fig_3_9_dfd_p6.png", "ภาพที่ 3.9 DFD ระดับที่ 1 ของกระบวนการที่ 6 ตรวจสอบและยืนยันผล"),
        ("fig_3_10_dfd_p7.png", "ภาพที่ 3.10 DFD ระดับที่ 1 ของกระบวนการที่ 7 แสดงผลการแข่งขันสาธารณะ"),
    ]):
        figure(doc, os.path.join(DIAGRAMS, fname), caption, width=5.6)

    # 3.2
    para(doc, "3.2 การออกแบบฐานข้อมูล (Database Design)", bold=True, space_after=6)

    para(doc, "3.2.1 แผนภาพความสัมพันธ์ของข้อมูล (Entity Relationship Diagram)",
         bold=True, indent=1.25, space_after=6)
    para(doc, "ระบบใช้ฐานข้อมูลเชิงสัมพันธ์ (MariaDB) ที่ผ่านการทำ Normalization ประกอบด้วยตาราง "
              "9 ตาราง โดยมีโครงสร้างความสัมพันธ์หลักแบบลำดับชั้น คือ การแข่งขัน (tournaments) "
              "หนึ่งรายการมีได้หลายรุ่น (divisions) แต่ละรุ่นประกอบด้วยหลายรอบ (rounds) แต่ละรอบมี"
              "หลายคู่แข่งขัน (matches) และแต่ละคู่แข่งขันประกอบด้วยหลายเกมตามรูปแบบ Best of "
              "(match_games) ส่วนทีม (teams) สังกัดรุ่นและโรงเรียน (schools) โดยมีสมาชิกทีม "
              "(team_members) ทีมละ 5 คน ทั้งนี้ตารางคู่แข่งขันมีความสัมพันธ์อ้างอิงตนเอง "
              "(next_match_id) เพื่อระบุคู่แข่งขันถัดไปที่ทีมชนะจะเลื่อนเข้าสู่รอบต่อไป และตารางเกม"
              "การแข่งขันอ้างอิงผู้ใช้งาน (users) ทั้งผู้อัปโหลดภาพหลักฐานและผู้ยืนยันผล "
              "โดยการแข่งขันแต่ละรายการระบุเอกลักษณ์ด้วยคู่ของปีและครั้งที่ (year, season) "
              "เพื่อรองรับการจัดการแข่งขันมากกว่าหนึ่งครั้งภายในปีเดียวกัน ดังภาพที่ 3.11",
         indent=1.25, space_after=6)
    figure(doc, os.path.join(DIAGRAMS, "fig_3_4_er.png"),
           "ภาพที่ 3.11 แผนภาพความสัมพันธ์ของข้อมูล (ER Diagram)", width=5.9)

    para(doc, "3.2.2 พจนานุกรมข้อมูล (Data Dictionary)", bold=True, indent=1.25, space_after=6)
    para(doc, "พจนานุกรมข้อมูลอธิบายรายละเอียดโครงสร้างของตารางในฐานข้อมูลทั้ง 9 ตาราง "
              "ประกอบด้วยชื่อฟิลด์ ชนิดข้อมูล คำอธิบาย และประเภทคีย์ โดย PK หมายถึงคีย์หลัก "
              "(Primary Key) FK หมายถึงคีย์นอก (Foreign Key) และ UQ หมายถึงฟิลด์ที่กำหนดค่าห้ามซ้ำ "
              "(Unique) ทั้งนี้กรณีที่หลายฟิลด์ในตารางเดียวกันกำกับ UQ ร่วมกัน หมายถึงการห้ามซ้ำ"
              "แบบผสม (Composite Unique) เช่น ฟิลด์ year และ season ของตารางการแข่งขัน "
              "ดังตารางที่ 3.2 ถึงตารางที่ 3.10",
         indent=1.25, space_after=8)
    for caption, rows in DD:
        data_table(doc, caption, rows)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
